"""Deterministic .docx assembly for Russian procedural documents.

Stage 2 of the drafting tool: takes a list of typed blocks
``[{"type": "...", "text": "..."}]`` and renders a fully formatted .docx with
the court-document layout baked in (A4, margins 3/1.5/2/2 cm, Times New Roman
14pt, 1.5 line spacing, justified body, 1.25 cm first-line indent, page numbers
top-center from page 2).

Ported from the standalone ``build_docx.py`` skill engine; the only changes are
that ``render_docx`` returns ``bytes`` (via an in-memory buffer) instead of
writing to disk, and the CLI wrapper is dropped. Depends only on ``python-docx``
(already a backend dependency), so it runs in-process — no sandbox/CLI needed.

Block types (``type``):
    header      — строка шапки (суд, дело, истец/ответчик, адрес). Блоком в
                  правой части листа (общий отступ слева ~8,25 см, строки по
                  левому краю — так требует ГОСТ Р 7.0.97-2016 для адресата).
    title       — заголовок документа. По центру, жирный.
    subtitle    — подзаголовок под title. По центру, жирный.
    h1          — заголовок раздела (I., II., 1.1.). Жирный, слева, отступ сверху.
    body        — обычный абзац. По ширине, красная строка 1.25.
    quote       — цитата практики. Как body, принудительно в «ёлочках».
    proshu      — абзац "...ПРОШУ:" (слово ПРОШУ — жирным).
    item        — пункт требования/списка. Как body.
    annex_h     — слово "Приложения:" (жирный, слева).
    annex_item  — пункт приложения. Как body.
    sign_left   — строка подписи/даты слева. Если сразу за ней идёт sign_right,
                  обе строки верстаются ОДНИМ абзацем: дата слева, подпись
                  справа через правый таб-стоп.
    sign_right  — строка подписи справа.
    spacer      — пустой абзац (пустая строка-отступ).
"""
from __future__ import annotations

import re
from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
FIRST_LINE = Cm(1.25)
LINE_SPACING = 1.5
# Ширина текстовой колонки: 21 см − поля 3 + 1,5 см. Правый таб-стоп подписи.
TEXT_WIDTH = Cm(16.5)
# Отступ блока шапки: правая половина листа, строки по левому краю.
HEADER_INDENT = Cm(8.25)

NBSP = " "  # неразрывный пробел
NDASH = "–"
MDASH = "—"

# Сокращения, после которых пробел перед числом должен быть неразрывным.
ABBR_BEFORE_NUM = [
    "№", "§", "ст", "стт", "п", "пп", "абз", "ч", "гл", "разд", "р", "л",
    "д", "дд", "т", "кв", "оф", "пом", "стр", "г", "гг", "обл", "ул", "пл",
]
# Короткие слова-предлоги, которые "тянут" следующий токен.
SHORT_WORDS = ["в", "во", "и", "к", "с", "со", "о", "об", "от", "до", "по",
               "на", "за", "из", "у", "не", "ни", "а", "но"]


def normalize_text(t: str) -> str:
    """Текстовая гигиена: кавычки, тире, неразрывные пробелы, лишние пробелы."""
    if not t:
        return t
    # 1. Убираем мягкие переносы и табы, схлопываем пробелы.
    t = t.replace("­", "").replace("\t", " ")
    t = re.sub(r"[  ]{2,}", " ", t)
    t = t.strip()

    # 2. Прямые кавычки -> ёлочки (парная расстановка).
    t = re.sub(r'"([^"]*)"', lambda m: "«" + m.group(1) + "»", t)
    t = t.replace('"', "")

    # 3. Дефис между пробелами -> длинное тире.
    t = re.sub(r"\s[-–]\s", f"{NBSP}{MDASH} ", t)

    # 4. Пробел перед знаком препинания убрать.
    t = re.sub(r"\s+([,.;:!?])", r"\1", t)

    # 5. Неразрывный пробел после сокращений перед числом: "ст. 196" -> "ст. 196".
    abbr = "|".join(sorted(ABBR_BEFORE_NUM, key=len, reverse=True))
    t = re.sub(rf"\b({abbr})\.\s+(?=[\dIVXЛ№])", rf"\1.{NBSP}", t, flags=re.IGNORECASE)
    t = re.sub(r"([№§])\s+", rf"\1{NBSP}", t)

    # 6. Инициалы: "О. В. Пашкин" и "Пашкин О. В." — неразрывно.
    t = re.sub(r"\b([А-ЯЁ])\.\s*([А-ЯЁ])\.\s*([А-ЯЁ][а-яё]+)",
               rf"\1.{NBSP}\2.{NBSP}\3", t)
    t = re.sub(r"\b([А-ЯЁ][а-яё]+)\s+([А-ЯЁ])\.\s*([А-ЯЁ])\.",
               rf"\1{NBSP}\2.{NBSP}\3.", t)

    # 7. Число + единица/слово, которые нельзя отрывать: "2026 г.", "100 руб.".
    t = re.sub(r"(\d)\s+(г|гг|руб|коп|млн|млрд|тыс|руб\.|%)\b", rf"\1{NBSP}\2", t)

    # 7а. Разряды чисел: "1 250 000" — неразрывно внутри числа.
    t = re.sub(r"(?<=\d) (?=\d{3}(?!\d))", NBSP, t)

    # 8. "ГК РФ", "ГПК РФ" и т.п. — неразрывно.
    t = re.sub(r"\b(ГК|ГПК|АПК|УК|УПК|КоАП|НК|ТК|СК)\s+РФ\b", rf"\1{NBSP}РФ", t)

    # 9. Короткие слова-предлоги приклеиваем к следующему слову.
    sw = "|".join(SHORT_WORDS)
    t = re.sub(rf"(^|\s)({sw})\s+", lambda m: f"{m.group(1)}{m.group(2)}{NBSP}", t,
               flags=re.IGNORECASE)

    t = re.sub(r" {2,}", " ", t)
    return t


def _set_run_font(run, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    # Критично для кириллицы — задать шрифт во всех слотах rFonts.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def _base_format(p, align, first_line: bool = True, space_before: int = 0,
                 space_after: int = 0, left_indent=None) -> None:
    pf = p.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.first_line_indent = FIRST_LINE if first_line else None
    pf.left_indent = left_indent


def _add_paragraph(doc, text, align, bold: bool = False, first_line: bool = True,
                   space_before: int = 0, space_after: int = 0, bold_word=None,
                   left_indent=None):
    p = doc.add_paragraph()
    _base_format(p, align, first_line, space_before, space_after, left_indent)
    text = normalize_text(text)
    if bold_word and bold_word in text:
        # Делим абзац, чтобы выделить одно слово (например ПРОШУ).
        idx = text.index(bold_word)
        before, word, after = text[:idx], bold_word, text[idx + len(bold_word):]
        if before:
            _set_run_font(p.add_run(before), bold)
        _set_run_font(p.add_run(word), True)
        if after:
            _set_run_font(p.add_run(after), bold)
    else:
        _set_run_font(p.add_run(text), bold)
    return p


def _setup_page(doc) -> None:
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    st = doc.styles["Normal"]
    st.font.name = FONT_NAME
    st.font.size = FONT_SIZE
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    st.element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)


def _set_core_properties(doc) -> None:
    # Дефолтный шаблон python-docx несёт автора "python-docx" и дату создания
    # 2013-12-23 — эти метаданные видны получателю через «Файл → Свойства».
    cp = doc.core_properties
    now = datetime.now(timezone.utc)
    cp.author = ""
    cp.last_modified_by = ""
    cp.comments = ""
    cp.created = now
    cp.modified = now


def _add_page_numbers(doc) -> None:
    """Номер страницы по центру верхнего поля, начиная со второй страницы
    (первая — без номера, как принято для процессуальных документов)."""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    p = sec.header.paragraphs[0]
    p.paragraph_format.alignment = _C
    for tag, attrs, text in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, "PAGE"),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        run = p.add_run()
        _set_run_font(run)
        el = OxmlElement(tag)
        for key, val in attrs.items():
            el.set(qn(key), val)
        if text:
            el.text = text
        run._element.append(el)


_J = WD_ALIGN_PARAGRAPH.JUSTIFY
_C = WD_ALIGN_PARAGRAPH.CENTER
_L = WD_ALIGN_PARAGRAPH.LEFT
_R = WD_ALIGN_PARAGRAPH.RIGHT


def _add_table(doc, rows: list[dict]) -> None:
    """Render a bordered Word table from ``[{"cells": [...]}, ...]``.

    First row is treated as a bold header. Ragged rows are padded to the widest
    row so python-docx never indexes out of range. Cells use the document font
    and single spacing for compactness; text goes through ``normalize_text``.
    """
    matrix = [
        list(r.get("cells", []) or [])
        for r in rows
        if isinstance(r, dict)
    ]
    if not matrix:
        return
    cols = max((len(r) for r in matrix), default=0)
    if cols == 0:
        return

    table = doc.add_table(rows=len(matrix), cols=cols)
    table.style = "Table Grid"  # single-line borders on every cell
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(matrix):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            # str(): в сохранённых blocks ячейки могут быть числами/None (JSON).
            raw = row[c_idx] if c_idx < len(row) else ""
            txt = normalize_text("" if raw is None else str(raw))
            para = cell.paragraphs[0]
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = None
            _set_run_font(para.add_run(txt), bold=(r_idx == 0))


def _add_signature_line(doc, left: str, right: str) -> None:
    """Дата слева и подпись справа ОДНОЙ строкой через правый таб-стоп."""
    p = doc.add_paragraph()
    _base_format(p, _L, first_line=False)
    p.paragraph_format.tab_stops.add_tab_stop(TEXT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    _set_run_font(p.add_run(normalize_text(left)))
    _set_run_font(p.add_run("\t"))
    _set_run_font(p.add_run(normalize_text(right)))


def _block_text(b: dict) -> str:
    txt = b.get("text", "")
    return txt if isinstance(txt, str) else str(txt or "")


def render_docx(blocks: list[dict]) -> bytes:
    """Render a list of typed blocks into a .docx and return its bytes."""
    doc = Document()
    _setup_page(doc)
    _set_core_properties(doc)
    _add_page_numbers(doc)
    blocks = [b for b in blocks if isinstance(b, dict)]
    i = 0
    while i < len(blocks):
        b = blocks[i]
        t = b.get("type", "body")
        txt = _block_text(b)
        if t == "header":
            _add_paragraph(doc, txt, _L, bold=False, first_line=False,
                           left_indent=HEADER_INDENT)
        elif t == "title":
            _add_paragraph(doc, txt, _C, bold=True, first_line=False,
                           space_before=6, space_after=0)
        elif t == "subtitle":
            _add_paragraph(doc, txt, _C, bold=True, first_line=False)
        elif t == "h1":
            _add_paragraph(doc, txt, _L, bold=True, first_line=False,
                           space_before=12, space_after=6)
        elif t == "quote":
            q = normalize_text(txt)
            # Оборачиваем в «ёлочки» только текст вовсе без кавычек — иначе
            # получаются вложенные/несбалансированные кавычки.
            if "«" not in q and "»" not in q:
                q = "«" + q + "»"
            _add_paragraph(doc, q, _J, bold=False, first_line=True)
        elif t == "proshu":
            _add_paragraph(doc, txt, _J, bold=False, first_line=True, bold_word="ПРОШУ")
        elif t == "item":
            _add_paragraph(doc, txt, _J, bold=False, first_line=True)
        elif t == "table":
            _add_table(doc, b.get("rows") or [])
        elif t == "annex_h":
            _add_paragraph(doc, txt, _L, bold=True, first_line=False, space_before=12)
        elif t == "annex_item":
            _add_paragraph(doc, txt, _J, bold=False, first_line=True)
        elif t == "sign_left":
            nxt = blocks[i + 1] if i + 1 < len(blocks) else None
            if nxt is not None and nxt.get("type") == "sign_right":
                _add_signature_line(doc, txt, _block_text(nxt))
                i += 2
                continue
            _add_paragraph(doc, txt, _L, bold=False, first_line=False)
        elif t == "sign_right":
            _add_paragraph(doc, txt, _R, bold=False, first_line=False)
        elif t == "spacer":
            p = doc.add_paragraph()
            _base_format(p, _J, first_line=False)
            _set_run_font(p.add_run(""))
        else:  # body (и любой неизвестный тип)
            _add_paragraph(doc, txt, _J, bold=False, first_line=True)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
