"""Document text extraction dispatcher.

Routes an uploaded document to text. Office/plain formats are parsed natively
off the FastAPI event loop (CPU-bound / subprocess). Everything OCR-able — PDFs
and images — is handed to the configured :class:`Recognizer`
(see ``app/rag_core/recognizers``). For PDFs the LLM recognizer first runs the
per-page text-layer gate (``recognizers/pdf_gate.py``): pages whose text layer
explains the rendered ink are taken from the layer directly; only pages that
actually need it (scans, garbled layers) are OCR'd. Native office parsers fall
back to the same recognizer when they yield nothing.
"""
from __future__ import annotations

import asyncio
import io
import logging
import os
import re
import tempfile
from contextvars import ContextVar
from dataclasses import dataclass
from typing import TYPE_CHECKING, Literal

if TYPE_CHECKING:  # avoid an import cycle: recognizers import helpers from here.
    from app.rag_core.recognizers.base import Recognizer

logger = logging.getLogger(__name__)

# Per-page scanned-PDF OCR knobs (consumed by the LLM recognizer).
PDF_PAGE_CONCURRENCY = 10
PDF_PAGE_MAX_RETRIES = 2
DOCUMENT_EXTRACTION_MAX_TOKENS = 16384

# Stage-A resilience knobs (see the OCR-resilience plan). The hard wall bounds
# ONE page (all retries included) so a stalled provider can't compound
# 300s-timeouts into hours; the budget bounds the whole extraction well under
# the serverless container kill; the breaker stops feeding a dying provider.
PDF_PAGE_HARD_TIMEOUT = 90.0        # s per page, retries included
PDF_EXTRACTION_HARD_BUDGET = 1500.0  # s from endpoint entry (download included)
PDF_DEADLINE_RESERVE = 120.0         # s kept back for assembly + DB insert
PDF_CIRCUIT_BREAKER_THRESHOLD = 20   # consecutive same-class page failures

# Absolute extraction deadline (monotonic loop time), set at endpoint entry so
# the S3 download and PDF parsing count against the budget too — a fixed
# wait-timeout deeper down would silently drift past the container kill.
_EXTRACTION_DEADLINE: ContextVar[float | None] = ContextVar(
    "extraction_deadline", default=None
)


def begin_extraction_deadline() -> float:
    """Start (or return the already-started) absolute deadline for this request."""
    deadline = _EXTRACTION_DEADLINE.get()
    if deadline is None:
        deadline = asyncio.get_running_loop().time() + PDF_EXTRACTION_HARD_BUDGET
        _EXTRACTION_DEADLINE.set(deadline)
    return deadline


def extraction_deadline() -> float | None:
    return _EXTRACTION_DEADLINE.get()

Strategy = Literal[
    "text", "docx", "doc", "vision", "llm-file", "pdf-pages", "pdf-text-layer", "sotaocr"
]

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".heic"}
_MIME_BY_EXT = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".txt": "text/plain",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
}


@dataclass
class ExtractedDocument:
    text: str
    raw_text_length: int
    truncated: bool
    strategy: Strategy
    # Per-page coverage (PDF per-page path only; None elsewhere). truncated=True
    # with pages_recognized < pages_total marks a partial extraction that the
    # endpoint keeps re-extractable instead of caching forever.
    pages_total: int | None = None
    pages_recognized: int | None = None


# --- mime / extension detection (mirror document-processing.ts) ---

def file_extension(filename: str) -> str:
    dot = filename.rfind(".")
    return filename[dot:].lower() if dot >= 0 else ""


def is_plain_text(mime_type: str, ext: str) -> bool:
    return mime_type.startswith("text/") or ext in {".txt", ".md", ".csv", ".json"}


def is_docx(mime_type: str, ext: str) -> bool:
    return (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or ext == ".docx"
    )


def is_doc(mime_type: str, ext: str) -> bool:
    return mime_type == "application/msword" or ext == ".doc"


def is_pdf(mime_type: str, ext: str) -> bool:
    return mime_type == "application/pdf" or ext == ".pdf"


def is_image(mime_type: str, ext: str) -> bool:
    return mime_type.startswith("image/") or ext in _IMAGE_EXTS


def is_image_ext(ext: str) -> bool:
    return ext.lower() in _IMAGE_EXTS


def mime_from_extension(ext: str) -> str:
    return _MIME_BY_EXT.get(ext.lower(), "application/octet-stream")


def normalize_result(
    raw_text: str,
    strategy: Strategy,
    *,
    truncated: bool = False,
    pages_total: int | None = None,
    pages_recognized: int | None = None,
) -> ExtractedDocument:
    cleaned = raw_text.replace("\x00", "").strip()
    return ExtractedDocument(
        text=cleaned,
        raw_text_length=len(cleaned),
        truncated=truncated,
        strategy=strategy,
        pages_total=pages_total,
        pages_recognized=pages_recognized,
    )


# --- native parsers (CPU-bound / subprocess -> run off the event loop) ---
# Note: PDFs are not parsed here — the recognizer's per-page gate (pdf_gate.py)
# decides text-layer vs OCR per page, so mixed/partial scans are still caught.

def _extract_docx_sync(data: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception:
        logger.warning("DOCX native parse failed; will fall back", exc_info=True)
        return ""


async def extract_docx(data: bytes) -> str:
    return await asyncio.to_thread(_extract_docx_sync, data)


def _extract_rtf_sync(data: bytes) -> str:
    """RTF -> text. Many ``.doc`` files are actually RTF (``{\\rtf`` magic), which
    antiword can't read. Codepage is taken from ``\\ansicpg`` (default cp1251 for
    Russian docs)."""
    try:
        from striprtf.striprtf import rtf_to_text

        m = re.search(rb"\\ansicpg(\d+)", data[:512])
        encoding = f"cp{m.group(1).decode()}" if m else "cp1251"
        # latin-1 keeps the RTF bytes 1:1; rtf_to_text decodes the \'xx escapes
        # itself using `encoding`.
        return rtf_to_text(data.decode("latin-1"), encoding=encoding, errors="ignore").strip()
    except Exception:
        logger.warning("RTF parse failed; will fall back", exc_info=True)
        return ""


async def extract_doc(data: bytes) -> str:
    """Legacy .doc -> text. Sniffs the real format (extension is unreliable):
    RTF (mislabeled .doc) via striprtf, a zip via the .docx parser, else the
    binary OLE .doc via antiword in a separate process (never blocks the loop)."""
    head = data.lstrip()[:5]
    if head == b"{\\rtf":
        return await asyncio.to_thread(_extract_rtf_sync, data)
    if data[:4] == b"PK\x03\x04":  # actually a .docx (OOXML zip) renamed .doc
        return await extract_docx(data)

    tmp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        proc = await asyncio.create_subprocess_exec(
            "antiword", tmp_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        out, err = await proc.communicate()
        if proc.returncode != 0:
            logger.warning(
                "antiword failed rc=%s: %s",
                proc.returncode, err.decode("utf-8", "replace")[:300],
            )
            return ""
        return out.decode("utf-8", "replace").strip()
    except FileNotFoundError:
        logger.warning("antiword not installed; .doc native parse unavailable")
        return ""
    except Exception:
        logger.warning("DOC native parse failed; will fall back", exc_info=True)
        return ""
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


# --- dispatcher ---

async def _recognize(
    recognizer: "Recognizer", data: bytes, mime_type: str, filename: str
) -> ExtractedDocument:
    result = await recognizer.recognize(data, mime_type, filename)
    return normalize_result(
        result.text,
        result.strategy,
        truncated=result.truncated,
        pages_total=result.pages_total,
        pages_recognized=result.pages_recognized,
    )


async def extract_text_from_document(
    data: bytes, mime_type: str, filename: str, recognizer: "Recognizer"
) -> ExtractedDocument:
    ext = file_extension(filename)

    if is_plain_text(mime_type, ext):
        # Match Node's buffer.toString('utf-8') (lossy replacement on bad bytes).
        return normalize_result(data.decode("utf-8", "replace"), "text")

    if is_docx(mime_type, ext):
        docx_text = await extract_docx(data)
        if docx_text:
            return normalize_result(docx_text, "docx")
        # SotaOCR can't read .docx -> it skips itself and the LLM recognizer's
        # file-attachment path handles it.
        return await _recognize(recognizer, data, mime_type, filename)

    if is_doc(mime_type, ext):
        doc_text = await extract_doc(data)
        if doc_text:
            return normalize_result(doc_text, "doc")
        return await _recognize(recognizer, data, mime_type, filename)

    # PDFs, images, and anything else -> recognizer (PDFs go through the
    # per-page text-layer gate inside the LLM recognizer).
    return await _recognize(recognizer, data, mime_type, filename)
