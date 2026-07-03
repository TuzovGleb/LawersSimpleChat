"""Layout tests for the deterministic .docx renderer.

These pin the court-document conventions the renderer must produce: header
block in the RIGHT half of the sheet, 14pt Times New Roman, page numbers,
date-left/signature-right on one line, clean core properties.
"""
import io
import zipfile
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Cm, Pt

from app.services.docx_builder import (
    HEADER_INDENT,
    TEXT_WIDTH,
    normalize_text,
    render_docx,
)


def _load(blocks):
    return Document(io.BytesIO(render_docx(blocks)))


def test_header_block_sits_in_right_half():
    doc = _load([
        {"type": "header", "text": "В Тверской районный суд г. Москвы"},
        {"type": "header", "text": "Истец: Иванов Иван Иванович"},
    ])
    assert HEADER_INDENT == Cm(8.25)
    for p in doc.paragraphs:
        # .twips: в docx значение хранится в twips, EMU-сравнение ловит округление
        assert p.paragraph_format.left_indent.twips == HEADER_INDENT.twips
        # Внутри правого блока строки ровняются по ЛЕВОМУ краю (ГОСТ), не вправо.
        assert p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert p.paragraph_format.first_line_indent is None


def test_normal_style_is_times_new_roman_14():
    doc = _load([{"type": "body", "text": "Абзац."}])
    st = doc.styles["Normal"]
    assert st.font.name == "Times New Roman"
    assert st.font.size == Pt(14)


def test_signature_pair_renders_on_one_line_with_right_tab():
    doc = _load([
        {"type": "sign_left", "text": "«10» июня 2026 г."},
        {"type": "sign_right", "text": "___________ /И. И. Иванов/"},
    ])
    assert len(doc.paragraphs) == 1
    p = doc.paragraphs[0]
    assert "\t" in p.text
    assert "2026" in p.text and "Иванов" in p.text
    stops = list(p.paragraph_format.tab_stops)
    assert len(stops) == 1
    assert stops[0].position.twips == TEXT_WIDTH.twips
    assert stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT


def test_lone_sign_left_stays_a_separate_paragraph():
    doc = _load([{"type": "sign_left", "text": "«10» июня 2026 г."}])
    assert len(doc.paragraphs) == 1
    assert "\t" not in doc.paragraphs[0].text


def test_quote_is_not_double_wrapped():
    doc = _load([
        {"type": "quote", "text": "Суд указал: «срок исчисляется заново»."},
        {"type": "quote", "text": "срок исчисляется заново"},
    ])
    first, second = (p.text for p in doc.paragraphs)
    assert first.count("«") == 1 and first.count("»") == 1
    assert second.startswith("«") and second.endswith("»")


def test_core_properties_do_not_leak_generator():
    doc = _load([{"type": "body", "text": "Абзац."}])
    cp = doc.core_properties
    assert cp.author == ""
    assert cp.last_modified_by == ""
    # Дефолтный шаблон python-docx датирован 2013-12-23.
    assert cp.created is not None and cp.created.year >= datetime.now().year - 1


def test_page_number_field_in_header_not_on_first_page():
    data = render_docx([{"type": "body", "text": "Абзац."}])
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        names = z.namelist()
        headers = [n for n in names if n.startswith("word/header")]
        assert headers, "page-number header part is missing"
        combined = "".join(z.read(n).decode("utf-8") for n in headers)
        assert "PAGE" in combined
        assert "titlePg" in z.read("word/document.xml").decode("utf-8")


def test_render_survives_garbage_blocks():
    data = render_docx([
        {"type": "body", "text": None},
        {"type": "неизвестный", "text": "текст"},
        "не-словарь",
        {"type": "table", "rows": []},
        {"type": "table", "rows": [{"cells": [2025, 100.5, None]}, {"cells": ["итого", 0]}]},
        {"no_type_at_all": True},
    ])
    assert zipfile.ZipFile(io.BytesIO(data)).testzip() is None
    doc = Document(io.BytesIO(data))
    cells = [c.text for row in doc.tables[0].rows for c in row.cells]
    assert "2025" in cells and "0" in cells and "None" not in cells


def test_digit_groups_are_non_breaking():
    out = normalize_text("Цена иска: 1 250 000 руб.")
    assert "1 250 000 руб." in out
    # Год и день не склеиваются с соседними словами по этому правилу.
    assert normalize_text("решение от 15 июня") == "решение от 15 июня"
