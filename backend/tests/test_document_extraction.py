"""Tests for the document extraction dispatcher (routing + native parsers).

PDFs/images route straight to the recognizer (no native PDF parsing); office
formats parse natively and fall back to the recognizer when empty.
"""
import io
import shutil

import pytest

from app.rag_core.recognizers.base import RecognitionResult
from app.services import document_extraction as de


# --- fake recognizer that records that it was invoked ---

class FakeRecognizer:
    def __init__(self, text="RECOGNIZED", strategy="sotaocr"):
        self._result = RecognitionResult(text=text, strategy=strategy)
        self.calls = 0

    async def recognize(self, data, mime_type, filename):
        self.calls += 1
        return self._result


# --- detection helpers ---

def test_detection_helpers():
    assert de.is_plain_text("text/markdown", ".md")
    assert de.is_plain_text("application/octet-stream", ".csv")
    assert de.is_docx("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "")
    assert de.is_docx("application/octet-stream", ".docx")
    assert de.is_doc("application/msword", "")
    assert de.is_pdf("application/pdf", "")
    assert de.is_pdf("application/octet-stream", ".pdf")
    assert de.is_image("image/png", "")
    assert de.is_image("application/octet-stream", ".heic")
    assert not de.is_image("application/pdf", ".pdf")


def test_file_extension_and_mime():
    assert de.file_extension("Дело.PDF") == ".pdf"
    assert de.file_extension("noext") == ""
    assert de.mime_from_extension(".docx").endswith("wordprocessingml.document")
    assert de.mime_from_extension(".unknown") == "application/octet-stream"


def test_normalize_result_strips_nul_and_trims():
    r = de.normalize_result("  he\x00llo \n", "text")
    assert r.text == "hello"
    assert r.raw_text_length == len("hello")
    assert r.truncated is False
    assert r.strategy == "text"


# --- routing (native parsers monkeypatched for determinism) ---

@pytest.fixture
def patch_natives(monkeypatch):
    state = {"docx": "", "doc": ""}

    async def fake_docx(data):
        return state["docx"]

    async def fake_doc(data):
        return state["doc"]

    monkeypatch.setattr(de, "extract_docx", fake_docx)
    monkeypatch.setattr(de, "extract_doc", fake_doc)
    return state


async def test_route_plain_text_no_recognizer():
    rec = FakeRecognizer()
    r = await de.extract_text_from_document(b"hello world", "text/plain", "a.txt", rec)
    assert r.strategy == "text" and r.text == "hello world"
    assert rec.calls == 0


async def test_route_docx_native_then_fallback(patch_natives):
    rec = FakeRecognizer(strategy="llm-file", text="LLM-FILE")
    patch_natives["docx"] = "real docx body"
    r = await de.extract_text_from_document(b"x", "x", "a.docx", rec)
    assert r.strategy == "docx" and r.text == "real docx body"
    assert rec.calls == 0

    patch_natives["docx"] = ""  # empty -> recognizer fallback
    r = await de.extract_text_from_document(b"x", "x", "a.docx", rec)
    assert r.strategy == "llm-file" and r.text == "LLM-FILE"
    assert rec.calls == 1


async def test_route_doc_native_then_fallback(patch_natives):
    rec = FakeRecognizer(strategy="llm-file")
    patch_natives["doc"] = "legacy doc body"
    r = await de.extract_text_from_document(b"x", "application/msword", "a.doc", rec)
    assert r.strategy == "doc"
    assert rec.calls == 0

    patch_natives["doc"] = ""
    r = await de.extract_text_from_document(b"x", "application/msword", "a.doc", rec)
    assert r.strategy == "llm-file"
    assert rec.calls == 1


def test_extract_rtf_sync():
    rtf = b"{\\rtf1\\ansi\\ansicpg1251 Hello RTF world}"
    assert "Hello RTF world" in de._extract_rtf_sync(rtf)


async def test_route_doc_rtf_parses_locally():
    # A ".doc" that is actually RTF ({\rtf) must be parsed locally (striprtf),
    # never handed to the recognizer (antiword can't read RTF -> used to 500).
    rec = FakeRecognizer(strategy="llm-file")
    rtf = b"{\\rtf1\\ansi\\ansicpg1251 RTFMARKER hello}"
    r = await de.extract_text_from_document(rtf, "application/msword", "a.doc", rec)
    assert r.strategy == "doc" and "RTFMARKER" in r.text
    assert rec.calls == 0


async def test_route_doc_zip_is_docx():
    # A ".doc" that is actually a .docx (PK zip) routes to the docx parser.
    import io as _io

    from docx import Document

    doc = Document()
    doc.add_paragraph("ZIPDOCMARKER")
    buf = _io.BytesIO()
    doc.save(buf)
    rec = FakeRecognizer(strategy="llm-file")
    r = await de.extract_text_from_document(buf.getvalue(), "application/msword", "a.doc", rec)
    assert r.strategy == "doc" and "ZIPDOCMARKER" in r.text
    assert rec.calls == 0


async def test_route_pdf_always_recognizer(patch_natives):
    # No native PDF parsing: even a "real" PDF goes through the recognizer.
    rec = FakeRecognizer(strategy="sotaocr", text="OCR TEXT")
    r = await de.extract_text_from_document(b"%PDF-1.4 ...", "application/pdf", "a.pdf", rec)
    assert r.strategy == "sotaocr" and r.text == "OCR TEXT"
    assert rec.calls == 1


async def test_route_image_and_unknown(patch_natives):
    rec = FakeRecognizer(strategy="sotaocr")
    r = await de.extract_text_from_document(b"x", "image/png", "a.png", rec)
    assert r.strategy == "sotaocr" and rec.calls == 1

    rec = FakeRecognizer(strategy="llm-file")
    r = await de.extract_text_from_document(b"x", "application/zip", "a.zip", rec)
    assert r.strategy == "llm-file" and rec.calls == 1


# --- native parsers on real files ---

def test_extract_docx_sync_real():
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph DOCXMARKER")
    doc.add_paragraph("Second paragraph here")
    buf = io.BytesIO()
    doc.save(buf)
    extracted = de._extract_docx_sync(buf.getvalue())
    assert "DOCXMARKER" in extracted
    assert "Second paragraph" in extracted


@pytest.mark.skipif(shutil.which("antiword") is None, reason="antiword not installed")
async def test_extract_doc_missing_antiword_or_runs():
    # Smoke: with garbage bytes antiword should fail cleanly -> "" (no raise).
    assert await de.extract_doc(b"garbage") == ""
